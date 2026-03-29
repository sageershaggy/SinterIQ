from docx import Document
from docx.shared import Pt
from datetime import datetime

OUTPUT_FILE = 'SinterIQ_Technical_Database_Document.docx'

content = {
    'title': 'SinterIQ — Technical Database Document',
    'meta': f'Generated from server.ts schema on {datetime.utcnow().date().isoformat()}',
    'sections': [
        {
            'heading': '1) Architecture Overview',
            'paragraphs': [
                'Frontend: React + Vite (src/) with tabs for Companies, Contacts, Commissions, Research, FollowUps, Import, Settings, Tracking, Users.',
                'Backend: Express server in server.ts using SQLite via better-sqlite3 (sintertechnik.db).',
                'AI: LLM configuration stored in app_settings and accessed via server LLM helper functions.'
            ]
        },
        {
            'heading': '2) Database Schema (primary tables)',
            'paragraphs': [
                'companies — core company records and lead tracking',
                'contacts — people associated with companies',
                'activities — logged interactions and follow-ups',
                'orders — orders and commission tracking',
                'users — application users (seeded by app)',
                'app_settings — key/value settings (LLM config, etc.)',
                'notes — free-text notes attached to companies',
                'research_history — searches and saved research results'
            ]
        },
        {
            'heading': '3) Key table: companies (selected fields and types)',
            'paragraphs': [
                '- id: INTEGER PRIMARY KEY AUTOINCREMENT',
                '- company_name: TEXT NOT NULL',
                '- country: TEXT NOT NULL',
                '- address, city, region: TEXT',
                '- industry, company_type: TEXT NOT NULL',
                '- employee_count: INTEGER, revenue_eur: REAL',
                '- website, company_email, legal_form, duns_number: TEXT',
                '- lead_score: INTEGER, lead_status: TEXT',
                "- technical_fit: TEXT, qualification_notes: TEXT, assigned_to: TEXT",
                "- ai_qualified_at: DATETIME, website_score/social_score/buying_probability: INTEGER",
                "- tracking_level/tracking_status/next_tracking_date/tracking_notes: TEXT/DATETIME",
                "- created_at, updated_at: DATETIME"
            ]
        },
        {
            'heading': '4) Key table: contacts (selected fields and types)',
            'paragraphs': [
                '- id: INTEGER PK, company_id: INTEGER FK -> companies(id)',
                '- full_name: TEXT NOT NULL, job_title, department: TEXT',
                '- email, phone_direct, phone_mobile, linkedin_url: TEXT',
                '- contacted_via, interest_reason, ceramic_bearing_experience, attempted_solution: TEXT',
                '- is_verified: BOOLEAN, verified_date: DATETIME, is_primary: BOOLEAN',
                '- notes, created_at, updated_at'
            ]
        },
        {
            'heading': '5) Key table: activities (selected fields)',
            'paragraphs': [
                '- id: INTEGER PK, company_id: INTEGER FK, contact_id: INTEGER FK (optional)',
                '- activity_type: TEXT, activity_date: DATETIME, performed_by: TEXT',
                '- subject, details, outcome: TEXT',
                '- follow_up_date: DATETIME, follow_up_done: BOOLEAN',
                '- attachments, created_at'
            ]
        },
        {
            'heading': '6) Key table: orders (selected fields)',
            'paragraphs': [
                '- id: INTEGER PK, company_id: INTEGER FK, contact_id: INTEGER FK (optional)',
                '- order_reference, order_date: DATETIME, order_value_eur: REAL',
                '- product_type, is_hybrid: BOOLEAN',
                '- commission_rate, commission_eur, payment_received, payment_date',
                '- commission_paid, commission_paid_date, notes, created_at'
            ]
        },
        {
            'heading': '7) Other tables',
            'paragraphs': [
                '- users: id, full_name, email, role, is_active, notes, timestamps',
                "- app_settings: key/value storage for LLM provider, model, api keys",
                '- notes: company-scoped messages (author, message, type, created_at)',
                '- research_history: stores research runs, results_json and saved company mapping'
            ]
        },
        {
            'heading': '8) Relations & Integrity',
            'paragraphs': [
                '- contacts.company_id -> companies.id',
                '- activities.company_id -> companies.id, activities.contact_id -> contacts.id (optional)',
                '- orders.company_id -> companies.id, orders.contact_id -> contacts.id (optional)',
                '- app handles uniqueness checks: company duplicate checks by name/website and contact email uniqueness per company.'
            ]
        },
        {
            'heading': '9) API Endpoints (DB operations — highlights)',
            'paragraphs': [
                '- Companies: GET /api/companies, GET /api/companies/:id, POST /api/companies, PUT/PATCH/DELETE, POST /api/companies/merge',
                '- Contacts: GET /api/contacts, POST /api/contacts, PUT/DELETE, POST /api/contacts/enrich (LLM)',
                '- Activities: POST /api/activities, follow-ups list, snooze/done endpoints, recent activities',
                '- Orders: GET /api/orders',
                '- Settings: GET/PUT /api/settings/llm',
                '- Export: GET /api/export/customer-tracker'
            ]
        },
        {
            'heading': '10) Normalization, dedupe, and utilities',
            'paragraphs': [
                '- Normalizers: required vs optional strings, numeric parsing, boolean flags.',
                "- normalizeCompanyNameForMatch and normalizeWebsiteHost used for dedupe/matching on import and research.",
                "- Default tracking_level='WATCHLIST' and tracking_status='PENDING' enforced on startup."
            ]
        },
        {
            'heading': '11) Recommended migration notes',
            'paragraphs': [
                '- If moving to PostgreSQL/MySQL: add explicit FKs, ON DELETE CASCADE where appropriate, and unique indexes (e.g., lower(company_name), lower(website), (company_id, lower(email))).',
                '- Add indices: companies(company_name), companies(website), companies(assigned_to), companies(lead_status), contacts(company_id), activities(company_id).',
                '- Convert textual enums (lead_status, activity_type) to reference tables or check constraints for integrity.'
            ]
        },
        {
            'heading': '12) Gap checklist',
            'paragraphs': [
                '- Core CRM tables present and seeded users: ✓',
                '- LLM settings persisted: ✓',
                '- Duplicate/merge logic implemented in server: ✓',
                '- Follow-up queue query implemented: ✓',
                '- Research history persisted: ✓'
            ]
        },
        {
            'heading': '13) Next steps / Handover',
            'paragraphs': [
                '- Add DB indices on heavy query columns before production migration.',
                '- Add integration tests that exercise merge logic and contact uniqueness checks.',
                '- Optionally convert to Postgres and add migrations (eg. using knex or sequelize).'
            ]
        }
    ]
}


def build_doc(out_path):
    doc = Document()
    doc.add_heading(content['title'], level=1)
    doc.add_paragraph(content['meta'])

    for section in content['sections']:
        doc.add_heading(section['heading'], level=2)
        for para in section['paragraphs']:
            # Use simple bullets for lines that start with '-'
            if para.strip().startswith('-'):
                p = doc.add_paragraph(para.lstrip('- ').strip(), style='List Bullet')
            else:
                doc.add_paragraph(para)

    # Save file
    doc.save(out_path)
    print(f'Wrote {out_path}')


if __name__ == '__main__':
    build_doc(OUTPUT_FILE)
